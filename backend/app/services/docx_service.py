from pathlib import Path

from docx import Document


def extract_docx(path: Path) -> tuple[list[dict], dict[str, str]]:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    props = document.core_properties
    metadata = {
        "author": props.author or "",
        "title": props.title or "",
        "subject": props.subject or "",
        "created": props.created.isoformat() if props.created else "",
        "modified": props.modified.isoformat() if props.modified else "",
    }
    metadata = {key: value for key, value in metadata.items() if value}

    text = "\n\n".join(paragraphs)
    return [{"page_number": 1, "text": text}], metadata

