import csv
import tempfile
from pathlib import Path
from xml.sax.saxutils import escape


def export_pages_to_txt(pages: list[dict], title: str) -> Path:
    output = Path(tempfile.NamedTemporaryFile(delete=False, suffix=".txt").name)
    text = [title, ""]
    for page in pages:
        text.append(f"Page {page['page_number']}")
        text.append(page.get("text", ""))
        text.append("")
    output.write_text("\n".join(text), encoding="utf-8")
    return output


def export_pages_to_csv(pages: list[dict]) -> Path:
    output = Path(tempfile.NamedTemporaryFile(delete=False, suffix=".csv").name)
    with output.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["page_number", "text"])
        for page in pages:
            writer.writerow([page["page_number"], page.get("text", "")])
    return output


def export_pages_to_pdf(pages: list[dict], title: str) -> Path:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = Path(tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name)
    doc = SimpleDocTemplate(str(output), pagesize=letter, leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

    for page in pages:
        story.append(Paragraph(f"Page {page['page_number']}", styles["Heading2"]))
        text = escape(page.get("text") or "No extracted text on this page.").replace("\n", "<br/>")
        story.append(Paragraph(text[:8000], styles["BodyText"]))
        story.append(Spacer(1, 12))

    doc.build(story)
    return output


def export_pages_to_docx(pages: list[dict], title: str) -> Path:
    from docx import Document

    output = Path(tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name)
    document = Document()
    document.add_heading(title, level=1)

    for page in pages:
        document.add_heading(f"Page {page['page_number']}", level=2)
        document.add_paragraph(page.get("text") or "No extracted text on this page.")

    document.save(output)
    return output


def export_pages_to_xlsx(pages: list[dict], title: str) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    output = Path(tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Extracted Text"
    sheet.append(["Document", title])
    sheet.append([])
    sheet.append(["Page", "Text"])
    for cell in sheet[3]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="146C94")

    for page in pages:
        sheet.append([page["page_number"], page.get("text", "")])

    sheet.column_dimensions["A"].width = 12
    sheet.column_dimensions["B"].width = 100
    workbook.save(output)
    return output
